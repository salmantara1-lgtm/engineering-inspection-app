import streamlit as st
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import json
import os

# --- الإعدادات الأساسية للصفحة ---
st.set_page_config(page_title="نظام استلامات العظم - أدوار الحداثة", page_icon="🏗️", layout="wide")

# --- كود CSS لإجبار الواجهة على دعم اللغة العربية (RTL) بالكامل ---
st.markdown("""
<style>
    .stApp {
        direction: rtl;
    }
    p, h1, h2, h3, h4, h5, h6, .stMarkdown, .stText, span, div {
        text-align: right !important;
    }
    .stTextInput input, .stTextArea textarea, .stSelectbox div {
        text-align: right !important;
        direction: rtl;
    }
    [data-testid="column"] {
        direction: rtl;
        text-align: right;
    }
    button[data-baseweb="tab"] {
        direction: rtl;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# نظام قواعد البيانات (للمشاريع والترقيم التراكمي)
# ==========================================
DB_FILE = "projects_bone_works_db.json"

def load_db():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return {"projects": [], "counters": {}}
    return {"projects": [], "counters": {}}

def save_db(db):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(db, f, ensure_ascii=False, indent=4)

def get_next_visit_id(project):
    db = load_db()
    return db["counters"].get(project, 0) + 1

def update_visit_counter(project):
    db = load_db()
    if project not in db["projects"]:
        db["projects"].append(project)
    if project not in db["counters"]:
        db["counters"][project] = 0
    db["counters"][project] += 1
    save_db(db)

# ==========================================
# قاعدة بيانات نماذج أعمال العظم (تم المطابقة حرفياً مع الصور المرفقة)
# ==========================================
INSPECTIONS_DATA = {
    "استلام اعمال الحفر": [
        "مراجعة الميزانية الشبكية للأرض الطبيعية قبل الحفر.",
        "مراجعة إحداثيات حدود الحفر مع حدود المبنى (حدود رفرفة الخرسانة العادية او الاحلال ان وجد طبقاً لتقرير الجسات).",
        "مراجعة منسوب التأسيس مع اللوحات وعمق الحفر من تقرير التربة وكذلك مع اقرب روبير.",
        "مراجعة تطهير قاع وجوانب الحفر واستواء الحفر.",
        "التأكد من ان نوع تربة قاع الحفر مطابقة لتقرير التربة ولا توجد اختلافات."
    ],
    "استلام اعمال الاحلال": [
        "التأكد من منسوب طبقة الاحلال (منسوب التأسيس).",
        "التأكد من سمك طبقة الاحلال طبقا لتوصيات الاستشارى.",
        "التأكد من طريقة الدمك المستخدمة (ميكانيكى/استاتيكى/هزاز) طبقا للمواصفات.",
        "التأكد من الوزن المستخدم فى الدمك طبقاً لتعليمات الاستشاري.",
        "التأكد من نتيجة اختبار الدمك قبل البدء فى الطبقة التالية.",
        " التأكد من مطابقة طبقة الاحلال للعينة المعتمدة."
    ],
    "استلام نجارة القواعد العادية": [
        "التأكد من نظافة منسوب التأسيس مع تمام دمك السطح.",
        "رفع اركان القواعد بجهاز Total Station ومطابقتها مع الرسومات.",
        "مطابقة المحاور الانشائية وصحة توقيع زواياها.",
        "مراجعة ابعاد القواعد وارتفاعاتها وحدود جوانبها.",
        "التقفيل الجيد لجوانب القواعد مع بعضها وتسديد الفتحات بين الالواح.",
        "مراجعة التقويات والتأكد من اتمامها بطريقة صحيحة.",
        "التأكد من وضع شرب صب القواعد (منسوب الظهر) بميزان القامة."
    ],
    "استلام نجارة القواعد الخرسانية المسلحة والسملات": [
        "يتم رفع جوانب النجارة بجهاز التوتال استيشن في حالة استخدامه للقواعد والسملات .",
        "التأكد من تطابق محاور القواعد مع المحاور الانشائية الصحيحة (يتم وضع المحور على ظهر الخرسانة العادية).",
        "التاكد من التقفيل الجيد لجوانب القواعد مع بعضها وتسديد الفتحات بين الالواح والتاكد من استقامة الاتجاهات وكذلك رأسية اجناب القواعد والسملات وفي حالة عدم عمل فرشة عادية اسفل السملات يتم توفير فرشة مناسبة تحتها .",
        "مراجعة ابعاد واتفاعات كل من القواعد والسملات بعد شد النجارة وكذلك التاكد من راسية اجناب القواعد والسملات.",
        "مراجعة اماكن الفتحات ومسارات الصحى والكهرباء.",
        "مراجعة التقويات (الزراجين والشكالات) لمنع الحركة اثناء الصب."
       "لتأكد من وضع شرب صب القواعد (منسوب الظهر) بميزان القامة."
    ],
    "استلام حديد تسليح الاساسات": [
        "التأكد من اقطار واعداد حديد التسليح ومطابقتها للمخططات.",
        "التأكد من اطوال التراكب (الوصلات) واماكنها.",
        "مراجعة تربيط الحديد بالسلك المجلفن بشكل جيد.",
         "التاكد من تركيب كراسي للحديد العلوي للقواعد ان وجد.",
        "مراجعة تركيب البسكويت (الغطاء الخرسانى) بالسمك المطلوب.",     
         "التأكد من اماكن فتحات ومسارات الصحي والكهرباء واماكن تثبيت الجوايط ... الخ.",
         "التأكد من اماكن اشاير حديد الاعمدة وربطها بكانات وتكون مستمرة داخل القاعدة .",
        "التأكد من نظافة حديد التسليح من الصدأ او الزيوت."
    ],
    "استلام نجارة الاعمدة الخرسانية": [
        "مراجعة قطاعات الاعمدة وابعادها.",
        "التأكد من راسية الاعمدة باستخدام ميزان الخيط (البلبل).",
        "مراجعة اماكن الاعمدة ومحاورها.",
        "التأكد من متانة التقويات (الاحزمة والزراجين).",
        "التأكد من منسوب نهاية صب العمود."
    ],
    "استلام تسليح الاعمدة والحوائط": [
        "التأكد من نظافة حديد التسليح وعدم وجود صدأ.",
        "التأكد من عدد واقطار الاسياخ.",
        "التاكد من راسية حديد الاعمدة وافقية الكانات.",
        "مراجعة المسافات بين الكانات وتكثيفها حسب الكود.",
        "التأكد من اقفال الكانات وتبادلها يمينا ويسارا.",
         "التاكد من تركيب كراسي للحديد العلوي للقواعد ان وجد.",
         "التأكد من تركيب بسكوت خرساني بين جوانب القاعدة وحديد تسليح القواعد طبقاً للغطاء الخرساني المطلوب تصميمياً .",
        "مراجعة طول اشارة العمود (طول الوصلة للعلوي).",
    \
    ],
    "استلام نجارة الاسقف الخرسانية تحت السقف": [
        "مراجعة القوائم (العروق) والمسافات بينها.",
        "مراجعة اماكن وصل العروق مع بعضها (فى حالة الارتفاعات العالية).",
        "التأكد من متانة الشدة الخشبية وافقيتها.",
        "مراجعة التقويات (البراندات) فى الاتجاهين.",
        "نظافة السطح الخشبى قبل التسليح."
    ],
    "استلام نجارة الاسقف الخرسانية فوق السقف والتصريح بالصب": [
        "مراجعة ابعاد البواكى والكمرات.",
        "مراجعة سقوط الكمرات.",
        "مراجعة اماكن وابعاد الفتحات (مناور / سلالم / مصاعد).",
        "مراجعة افقية السطح باستخدام ميزان القامة.",
        "مراجعة التقويات للكمرات (نهايات السقف)."
    ],
    "استلام تسليح الخرسانة المسلحة": [
        "مراجعة حديد الفرش والغطاء للبلاطات.",
        "مراجعة حديد الكمرات (الساقط، المعلق، المكسح).",
        "التأكد من أماكن وأطوال وصلات التراكب في الحديد.",
        "مراجعة تسليح السلالم وأشيرها.",
        "مراجعة التمديدات الكهربائية والصحية داخل السقف قبل الصب."
    ],
    "استلام العزل للاساسات": [
        "التأكد من جفاف ونظافة الأسطح المراد عزلها.",
        "التأكد من عمل وزرة (رقبة زجاجة) في الزوايا.",
        "دهان طبقة الأساس (البرايمر) وتغطية السطح بالكامل.",
        "لحام لفائف العزل (الرولات) جيداً والتأكد من مسافة الركوب.",
        "حماية العزل بعد التنفيذ من أي أضرار."
    ],
    "فحص واستلام اعمال صب الخرسانة": [
        "التأكد من إجهاد الخرسانة الموردة وتطابقها مع التصميم.",
        "أخذ عينات المكعبات لاختبار الضغط وإجراء اختبار الهبوط.",
        "التأكد من استخدام الهزاز الميكانيكي بشكل منتظم وجيد.",
        "منع إضافة الماء للخلاطة في الموقع نهائياً.",
        "تسوية سطح الخرسانة وتأمين معالجتها بالماء بعد الجفاف."
    ],
    "استلام حدادة الحوائط": [
        "مراجعة سمك الحائط الخرساني وطوله حسب المخطط.",
        "التأكد من رأسية النجارة واستقامتها.",
        "مراجعة تسليح الحائط (الشبكة المزدوجة) والكانات الرابطة.",
        "التأكد من متانة الشدة والزراجين.",
        "تثبيت قطع البسكويت لضمان الغطاء الخرساني."
    ]
}

# ==========================================
# البنود الإدارية الثمانية (كما في النموذج المرجعي)
# ==========================================
ADMIN_CHECKS = [
    "التأكد من إعتماد المواد المستخدمه فى تنفيذ البند",
    "التأكد من إعتماد رسومات الورشه اللازمه لتنفيذ البند",
    "التأكد من إعتماد كيفية التنفيذ لهذا البند",
    "التأكد من إجراء الإختبارات اللازمه طبقا لمواصفة البند",
    "التأكد من إستلام البنود المؤسس عليها البند المطلوب إستلامه",
    "التأكد من مراجعة رسومات التصميم والمواصفات الفنيه",
    "التأكد من توافر شروط السلامه المهنيه لإستلام البند",
    "التأكد من إزاله المخلفات الناتجه عن تنفيذ البند"
]

# --- دوال تنسيق الوورد (RTL للغة العربية) ---
def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)

def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement('w:bidiVisual')
    bidiVisual.set(qn('w:val'), '1')
    tblPr.append(bidiVisual)

# --- دالة التقرير (Word) ---
def create_docx(data):
    doc = Document()
    
    header = doc.add_heading(f"نموذج {data['inspection_type']}", level=1)
    set_rtl(header)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # البيانات الأساسية
    h2 = doc.add_heading('البيانات الأساسية للمشروع:', level=2)
    set_rtl(h2)
    p = doc.add_paragraph()
    set_rtl(p)
    p.add_run(f"اسم المشروع: {data['project']}\n").bold = True
    p.add_run(f"الاستشاري: {data['consultant']}\n")
    p.add_run(f"ممثل الاستشاري: {data['engineer']}\n")
    p.add_run(f"ممثل المقاول: {data['contractor']}\n")
    p.add_run(f"رقم الزيارة (الاستلام): {data['inspection_no']}\n")
    p.add_run(f"تاريخ الفحص: {data['date']}\n")

    # جدول الإدارة
    h3 = doc.add_heading('مراجعة النقاط الإدارية:', level=2)
    set_rtl(h3)
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'
    set_table_rtl(t1) 
    t1.rows[0].cells[0].text = 'بند المراجعة'
    t1.rows[0].cells[1].text = 'الحالة'
    for key, value in data['general_checks'].items():
        row = t1.add_row().cells
        row[0].text = key
        row[1].text = "تم" if value else "غير مطبق"

    # جدول الفحص الفني (بناءً على طلبك: ملاحظات - النتيجة - بند الفحص)
    h3_tech = doc.add_heading('بنود المراجعة الفنية:', level=2)
    set_rtl(h3_tech)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    set_table_rtl(t2)
    t2.rows[0].cells[0].text = 'بند الفحص'
    t2.rows[0].cells[1].text = 'النتيجة'
    t2.rows[0].cells[2].text = 'الملاحظات'
    
    for key, value in data['tech_checks'].items():
        row = t2.add_row().cells
        row[0].text = key
        row[1].text = value['result']
        row[2].text = value['note']

    # التقييم
    h4 = doc.add_heading('القرار النهائي وملاحظات المهندس:', level=2)
    set_rtl(h4)
    p_final = doc.add_paragraph()
    set_rtl(p_final)
    p_final.add_run(f"حالة الاعتماد: {data['status']}\n").bold = True
    p_final.add_run(f"توجيهات عامة: {data['comments']}")

    # إدراج الصور
    if data['photos']:
        doc.add_page_break()
        h_img = doc.add_heading('التوثيق الفوتوغرافي للموقع:', level=1)
        set_rtl(h_img)
        for idx, photo in enumerate(data['photos']):
            try:
                doc.add_picture(BytesIO(photo['image'].getvalue()), width=Inches(4.5))
                p_img = doc.add_paragraph(f"صورة ({idx+1}): {photo['desc']}")
                set_rtl(p_img)
            except: pass

    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- واجهة الاستخدام (Main App UI) ---
st.title("🏗️ نظام استلامات العظم - أدوار الحداثة")
st.markdown("---")

# ==========================================
# 1. الواجهة الرئيسية: بيانات المشروع
# ==========================================
st.subheader("⚙️ أولاً: بيانات المشروع والاستلام (حقول إلزامية)")

db_data = load_db()

col1, col2 = st.columns(2)

with col1:
    proj_options = db_data["projects"] + ["+ إضافة مشروع جديد..."]
    selected_proj = st.selectbox("اختر المشروع:", proj_options)
    
    if selected_proj == "+ إضافة مشروع جديد...":
        project_name = st.text_input("أدخل اسم المشروع الجديد (مطلوب):")
    else:
        project_name = selected_proj
        
    consultant = st.text_input("الاستشاري", "شركة ادوار الحداثة للاستشارات الهندسية")
    engineer = st.text_input("ممثل الاستشاري", "م.سلمان تاره")

with col2:
    stage = st.selectbox("المرحلة المراد استلامها:", list(INSPECTIONS_DATA.keys()))
    contractor = st.text_input("ممثل المقاول (مطلوب):", placeholder="اكتب اسم ممثل المقاول هنا")
    date_val = st.date_input("تاريخ الفحص", datetime.now()).strftime('%Y-%m-%d')

next_visit_no = get_next_visit_id(project_name) if project_name else 1
st.success(f"🔢 بناءً على سجلات هذا المشروع، رقم هذه الزيارة سيكون: **{next_visit_no}**")

st.markdown("---")

# ==========================================
# 2. تبويبات الفحص الفني والصور
# ==========================================
st.subheader("🛠️ ثانياً: الفحص الهندسي الميداني")

tab1, tab2, tab3 = st.tabs(["📝 بنود المراجعة", "📸 إرفاق الصور (حتى 10)", "📤 الاعتماد والإصدار"])

with tab1:
    st.write("### البنود الإدارية (التحقق قبل الاستلام)")
    gen_checks = {}
    for admin_item in ADMIN_CHECKS:
        gen_checks[admin_item] = st.checkbox(admin_item, value=True)
    
    st.divider()
    st.write(f"### الفحص الفني: {stage}")
    tech_results = {}
    for item in INSPECTIONS_DATA[stage]:
        st.markdown(f"**🔹 {item}**")
        # الخيارات هنا مطابقة للأعمدة في الصور المرفقة تماماً
        res = st.radio("النتيجة:", ["مطابق", "غير مطابق"], horizontal=True, key=f"res_{item}", label_visibility="collapsed")
        note = st.text_input("ملاحظات عامة:", key=f"note_{item}", placeholder="أضف تعليقاً على هذا البند...")
        tech_results[item] = {"result": res, "note": note if note else "-"}
        st.markdown("---")

with tab2:
    st.write("### توثيق الموقع (اختياري - حتى 10 صور)")
    photos_list = []
    for i in range(1, 11):
        with st.expander(f"إرفاق صورة رقم {i}"):
            f = st.file_uploader(f"التقط أو ارفع صورة {i}", type=['jpg','jpeg','png'], key=f"img_{i}")
            d = st.text_input(f"وصف الصورة {i}", key=f"d_{i}")
            if f:
                st.image(f, use_column_width=True)
                photos_list.append({"image": f, "desc": d if d else "-"})

with tab3:
    st.write("### الاعتماد النهائي للتقرير")
    final_no = st.text_input("رقم الزيارة للتثبيت في التقرير:", value=str(next_visit_no))
    decision = st.radio("القرار النهائي للاستلام:", ["مقبول", "مقبول مع ملاحظات يتم تلافيها", "مرفوض ويعاد التقديم"])
    general_notes = st.text_area("توجيهات عامة لممثل المقاول:")

    st.markdown("---")
    if st.button("✅ تأكيد البيانات وإصدار التقرير", use_container_width=True):
        if not project_name or not project_name.strip() or not contractor or not contractor.strip():
            st.error("⚠️ يرجى التأكد من تعبئة جميع بيانات المشروع في الأعلى (اسم المشروع وممثل المقاول) فهي حقول إلزامية.")
        else:
            update_visit_counter(project_name)
            final_data = {
                "inspection_type": stage, "project": project_name, "consultant": consultant,
                "engineer": engineer, "contractor": contractor, "inspection_no": final_no,
                "date": date_val, "general_checks": gen_checks, "tech_checks": tech_results,
                "photos": photos_list, "status": decision, "comments": general_notes
            }
            report = create_docx(final_data)
            
            file_name = f"{stage} - {project_name}.docx"
            
            st.success(f"🎉 تم حفظ بيانات الزيارة بنجاح للمشروع: {project_name}")
            st.download_button("📥 تحميل التقرير (Word)", data=report, file_name=file_name, use_container_width=True)
